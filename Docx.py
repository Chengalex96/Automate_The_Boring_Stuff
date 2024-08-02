# Reading and Editing Word Documents

import docx

d = docx.Document('demo.docx')


# Paragaphs are a collection of runs
print(d.paragraphs)

print(d.paragraphs[0].text)
print(d.paragraphs[1].text)

p = d.paragraphs[1]

# A new run whenever there's a change in style

print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)

p.runs[3].underline = True # Made the worde underlined and added the word
p.runs[3].text = 'italic and underlined'

d.save('demo2.docx')

# On word: Ctrl + alt + shift + s to see a list of styles

p.style = 'Title' # changes title style, styles are set
d.save('demo3.docx') 

# Blank document

d = docx.Document()
d.add_paragraph('Hello, this is a paragraph.')
d.add_paragraph('ANother paragraph.')
d.save('demo4.docx') 

p = d.paragraphs[0]
p.add_run('This is a new run')
p.runs[1].bold = True # made the new line we just added as bolded

d.save('demo5.docx')  

# Can only modify by adding it to the end
print('New Project')

# A string of all the texts in the word docs without style
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
    
print(getText('demo.docx'))